import streamlit as st
import pdfplumber
from rembg import remove
from PIL import Image
from io import BytesIO
from docx import Document
from docx.shared import Cm
import st_file_uploader as stf

def extraer_texto_pdf(pdf_file):
    try:
        pdf_bytes = BytesIO(pdf_file.getvalue())
        pdf_bytes.seek(0)
        
        with pdfplumber.open(pdf_bytes) as pdf:
            texto = "\n".join([page.extract_text() or "" for page in pdf.pages])
        
        return texto if texto else "No se pudo extraer texto del PDF."
    except Exception as e:
        return f"Error al procesar el PDF: {str(e)}"

def main():
    st.title("Manejo de Archivos Subidos")
    
    st.header("Carga y procesamiento de PDF")
    with st.container():
        pdf_file = stf.file_uploader("Selecciona un archivo PDF", type=["pdf"])
    
        if pdf_file is not None:
            if st.button("Extraer texto del PDF"):
                texto = extraer_texto_pdf(pdf_file)
                st.text_area("Texto extraído:", texto, height=200)
    
    st.header("Insertar imagen y texto en documento Word")
    with st.container():
        docx_file = stf.file_uploader("Selecciona una plantilla de Word", type=["docx"])
        image_file = stf.file_uploader("Selecciona una imagen", type=["jpg", "jpeg", "png"])
    
    if docx_file is not None and pdf_file is not None and image_file is not None:
        if st.button("Insertar en documento"):
            try:
                with st.spinner("Procesando documento..."):
                    # Abrir documento Word
                    doc_bytes = BytesIO(docx_file.getvalue())
                    doc = Document(doc_bytes)
                    
                    # Extraer texto del PDF
                    texto_pdf = extraer_texto_pdf(pdf_file)
                    doc.add_paragraph(texto_pdf)
                    
                    # Procesar imagen con PIL y remover fondo con rembg
                    img_bytes = BytesIO(image_file.getvalue())
                    image = Image.open(img_bytes)
                    
                    # Elimina el fondo de la imagen usando rembg
                    image_sin_fondo = remove(image)
                    
                    # Guardar la imagen procesada en un objeto BytesIO en formato PNG
                    output_image = BytesIO()
                    image_sin_fondo.save(output_image, format="PNG")
                    output_image.seek(0)
                    
                    # Insertar imagen en el documento
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()
                    run.add_picture(output_image, height=Cm(5))
                    
                    # Guardar documento modificado
                    output_doc = BytesIO()
                    doc.save(output_doc)
                    output_doc.seek(0)
                    
                    st.download_button(
                        label="Descargar documento con imagen y texto",
                        data=output_doc,
                        file_name="documento_con_texto_imagen.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            except Exception as e:
                st.error(f"Error al procesar el documento: {e}")
    
    st.header("Ejemplos de uso del File Uploader personalizado")
    
    with st.container():
        st.subheader("Versión completamente personalizada")
        custom = stf.create_custom_uploader(
            uploader_msg="¡Suelta tu archivo aquí!",
            limit_msg="Tamaño máximo 200MB",
            button_msg="Seleccionar Archivo",
            icon="MdFileUpload"
        )
        file_custom = custom.file_uploader(
            "Subir con texto personalizado",
            type=["xlsx", "csv"],
            accept_multiple_files=True,
        )
    
    with st.container():
        st.subheader("Uso básico (Inglés por defecto)")
        file = stf.file_uploader("Subir un archivo CSV", type="csv")
    
    with st.container():
        st.subheader("Versión en Español")
        file_es = stf.es.file_uploader("Sube un archivo CSV", type="csv")
    
    with st.container():
        st.subheader("Francés con personalización")
        file_fr = stf.fr.file_uploader(
            "Télécharger un fichier",
            type=["jpg", "png", "gif"],
            accept_multiple_files=True,
            button_msg="Sélectionner une image",
        )
    
    with st.container():
        st.subheader("Demostración de múltiples tipos de archivo")
        file_types = stf.file_uploader(
            "Subir documentos",
            type=["csv", "pdf", "docx"],
            accept_multiple_files=True,
        )

if __name__ == "__main__":
    main()