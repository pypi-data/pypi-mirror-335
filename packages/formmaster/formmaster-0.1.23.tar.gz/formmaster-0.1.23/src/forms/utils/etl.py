from docx.api import Document
from glob import glob

import pandas as pd
import os
import string
import re

def extract_eng(s):
    s = s + ' '

    # normalise punctuations
    s = s.replace("’", "'")

    idx = len(s) - 1
    for i in reversed(range(len(s))):
        if s[i] not in list(string.printable):
            break
        idx = i
    text = s[idx:].strip()
    text = re.sub(r'\s+', ' ', text)
    return text

def get_school(a):
    def match_school(a, names):
        for n in names:
            if re.search(n, a):
                return True
        return False
    
    if match_school(a, ["悉尼大学", "Taylor's college", "USYD"]):
        return "USYD"
    elif match_school(a, ["昆士兰大学", "Queen’s college", "UQ"]):
        return "UQ"
    elif match_school(a, ["新南威尔士大学", "UNSW global", "UNSW"]):
        return "UNSW"
    elif match_school(a, ["莫纳什大学", "Monash"]):
        return "Monash"
    elif match_school(a, ["墨尔本大学", "UNIMELB"]):
        return "UNIMELB"
    elif match_school(a, ["国立大学", "ANU"]):
        return "ANU"
    else:
        return "Other"

def load_table(app_tbl = '/home/hmei/data/13. 懿心ONE Bonnie/0209 曾政源/曾政源-澳洲大学申请信息表2023.docx'):
    document = Document(f'{app_tbl}')
    data = {'Number': 0}
    for table in document.tables[:2]:
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)
            line = []
            for t in text:
                if len(line) > 0 and line[-1] == t:
                    continue
                else:
                    line.append(t)
            for i in range(len(line)//2):
                k = extract_eng(line[i*2])
                v = extract_eng(line[i*2+1])
                data[k] = v
                if k == 'Home Address':
                    addr_parts = v.split(',')
                    province = addr_parts[-1]
                    city = addr_parts[-2]
                    addr_parts_len = len(addr_parts)
                    if addr_parts_len >= 5:
                        line3 = addr_parts[-3]
                        line2 = addr_parts[-4]
                        line1 = ', '.join(addr_parts[:-4])
                    else:
                        line1 = addr_parts[0]
                        line2 = ''
                        if len(addr_parts) >= 4:
                            line2 = addr_parts[1]
                        line3 = ''
                    data.update({'province': province, 'city': city, 'line1': line1, 'line2': line2, 'line3': line3})

    df_personal = pd.DataFrame.from_dict({'keys': data.keys(), 'values': data.values()})

    lines = []
    for i, row in enumerate(document.tables[2].rows):
        text = list(cell.text for cell in row.cells[1:])
        text = [extract_eng(t) for t in text]
        lines.append(text)
    df_edu = pd.DataFrame(lines[1:], columns = lines[0])

    lines = []
    for i, row in enumerate(document.tables[3].rows):
        text = list(cell.text for cell in row.cells)
        #text = [extract_eng(t) for t in text]
        lines.append(text)

    df_application = pd.DataFrame(lines[1:], columns = [extract_eng(t) for t in lines[0]])
    df_application['Proposed School'] = df_application['Proposed School'].apply(lambda a: get_school(a))
    df_application['Proposed Course with Corresponding Links'] = \
        df_application['Proposed Course with Corresponding Links'].apply(lambda a: extract_eng(a[::-1])[::-1])
    df_application['CRICOS'] = df_application['CRICOS'].apply(lambda a: extract_eng(a))
    df_application['Commencement Date(mm/yyyy)'] = df_application['Commencement Date(mm/yyyy)'].apply(lambda a: extract_eng(a))

    return [data, df_edu, df_application, df_personal]

def load_ielts(app_file = '/home/hmei/data/13. 懿心ONE Bonnie/0209 曾政源/'):
    return

def load(_dir):
    students = []
    if os.path.isdir(_dir):
        files = glob(f'{_dir}/**', recursive=True)
    else:
        files = [_dir]

    for file in files:
        if re.search('Application Form for AU University', file) or re.search('澳洲大学申请信息表', file):
            print(file)
            try:
                student = load_table(file)
                students.append(student)
            except Exception as e:
                print(str(e))
                print(f'Failed to load file {file}.')


    
    return students

