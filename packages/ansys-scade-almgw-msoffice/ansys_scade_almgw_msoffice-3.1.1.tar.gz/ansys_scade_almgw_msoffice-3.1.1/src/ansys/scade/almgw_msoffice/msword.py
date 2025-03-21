# Copyright (C) 2024 - 2025 ANSYS, Inc. and/or its affiliates.
# SPDX-License-Identifier: MIT
#
#
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in all
# copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
# IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
# FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
# AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
# LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
# OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
# SOFTWARE.

"""Extraction of the requirements of a MS Word document (DOCX)."""

from pathlib import Path
from re import compile
from typing import Any, Dict

from docx import Document

from ansys.scade.pyalmgw.documents import ReqDocument, ReqProject, Requirement, Section

heading_regex = compile('Heading (\\d+)')


def read_heading(style: Any) -> int:
    """Return the heading level of a paragraph or 0."""
    if style.builtin or not style.base_style:
        res = heading_regex.search(style.name)
        return int(res.group(1)) if res else 0
    else:
        return read_heading(style.base_style)


class Parser:
    """Parser for MS Word documents (DOCX)."""

    def __init__(self):
        """Initialize the parser."""
        self.current_req = None
        # stack of sections
        self.sections = []
        # dictionary of requirements
        self.requirements = {}
        # invalid requirements ids count
        self.anonymous_req_count = 0

    def add_req_id(self, id: str):
        """Add a new requirement to the document."""
        if not id:
            self.anonymous_req_count += 1
            id = '<Missing Requirement ID>_%d' % self.anonymous_req_count
        # current section
        top = self.sections[-1]
        self.current_req = Requirement(top, id)
        self.requirements[id] = self.current_req

    def append_req_text(self, text: str):
        """Append text to the current requirement's description."""
        if self.current_req is None:
            # error in the input document: create an anonymous requirement
            self.add_req_id('')
            assert self.current_req
        if self.current_req.description:
            self.current_req.description += '\n'
        self.current_req.description += text

    def add_heading(self, level: int, text: str):
        """Add a new section to the document."""
        top = self.sections[-1]
        current_level = len(self.sections) - 1
        # find the parent of the new section
        while current_level >= level and len(self.sections) > 0:
            self.sections.pop()
            top = self.sections[-1]
            current_level = len(self.sections) - 1
        number = top.number if len(self.sections) > 1 else ''
        # add sections if needed
        while current_level + 1 < level:
            current_level += 1
            number += '1.'
            section = Section(top, number, '<Missing Section>')
            self.sections.append(section)
            top = section
        # add the new section
        number += '{0}.'.format(len(top.sections) + 1)
        section = Section(top, number, text)
        self.sections.append(section)
        self.current_req = None

    def parse(self, req_doc: ReqDocument, req_style: str, text_style: str):
        """Build the document model from a MS Word document."""
        document = Document(str(req_doc.path))

        # initialize the stack of sections with the document
        self.sections = [req_doc]

        for p in document.paragraphs:
            lvl = read_heading(p.style)
            txt = p.text.strip()
            if lvl > 0:
                self.add_heading(lvl, txt)
            elif p.style and p.style.name == req_style:
                self.add_req_id(txt)
            elif p.style and p.style.name == text_style:
                self.append_req_text(txt)
            else:
                # ignore the paragraph
                pass


def add_document(
    project: ReqProject, path: Path, req_style: str, text_style: str
) -> Dict[str, Requirement]:
    """Parse the input MS Word document and return the contained requirements hierarchy."""
    req_doc = ReqDocument(project, path.as_posix(), path.name)
    parser = Parser()
    parser.parse(req_doc, req_style, text_style)
    return parser.requirements
